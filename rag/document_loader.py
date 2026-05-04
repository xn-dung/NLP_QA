from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document as DocxDocument
from pypdf import PdfReader

from .config import SUPPORTED_EXTENSIONS


@dataclass(frozen=True)
class RawDocument:
    source: str
    text: str


def load_documents(folder: Path) -> list[RawDocument]:
    folder.mkdir(parents=True, exist_ok=True)
    docs: list[RawDocument] = []

    for path in sorted(folder.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        text = _read_file(path)
        if text.strip():
            docs.append(RawDocument(source=str(path.relative_to(folder)), text=text))

    return docs


def _read_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    return ""


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"\n[Page {page_number}]\n{text}")
    return _clean_text("\n".join(pages))


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return _clean_text("\n".join(paragraph.text for paragraph in doc.paragraphs))


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"([A-Za-z])-\s*\n\s*([A-Za-z])", r"\1\2", text)
    text = re.sub(r"(?<![.!?:;\]\)])\n(?!\[Page \d+\])", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
